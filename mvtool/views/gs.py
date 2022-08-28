# coding: utf-8
#
# Copyright (C) 2022 Helmar Hutschenreuter
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.

import re
import shutil
from tempfile import NamedTemporaryFile
from fastapi import APIRouter, Depends, Response, UploadFile
from fastapi_utils.cbv import cbv
import docx

from .projects import ProjectsView
from .requirements import RequirementsView
from ..database import CRUDOperations
from ..models import GSBaustein, Requirement
from .. import errors


class ParagraphsWrapper:
    def __init__(self, paragraphs):
        self.paragraphs = paragraphs
        self.index = 0

    def reset(self):
        self.index = 0

    def next(self):
        if self.index >= len(self.paragraphs):
            return None
        else:
            self.index += 1
            return self.paragraphs[self.index - 1]

    def previous(self):
        if self.index <= 1:
            return None
        else:
            self.index -= 1
            return self.paragraphs[self.index - 1]

    @property
    def current(self):
        if self.index > len(self.paragraphs):
            return None
        else:
            return self.paragraphs[self.index - 1]


class GSBausteinParser:
    _requirement_title_re_1 = re.compile(
        r"^\s*([A-Z]+(\.[0-9]+)+\.A[0-9]+)\s*(.*?)\s*(\[([^\]]*)\])?\s*\((B|S|H)\)\s*$"
    )
    _requirement_title_re_2 = re.compile(
        r"^\s*([A-Z]+(\.[0-9]+)+\.A[0-9]+)\s*(.*?)\s*\((B|S|H)\)\s*(\[([^\]]*)\])?\s*$"
    )
    _gs_baustein_title_re = re.compile(r"^\s*([A-Z]+(\.[0-9]+)+)\s*(.*?)\s*$")

    @classmethod
    def parse(cls, filename):
        try:
            word_document = docx.Document(filename)
        except docx.opc.exceptions.PackageNotFoundError:
            raise errors.ValueHttpError("Word file seems to be corrupted")
        paragraphs = ParagraphsWrapper(word_document.paragraphs)
        gs_baustein = cls._parse_baustein(paragraphs)
        return gs_baustein

    @classmethod
    def _parse_requirement_texts(cls, paragraphs, requirement):
        text = ""
        while paragraphs.next():
            if paragraphs.current.style.name == "Normal":
                text += paragraphs.current.text
            else:
                paragraphs.previous()
                break

        requirement.description = text
        return requirement

    @classmethod
    def _parse_requirement_title(cls, text):
        match = cls._requirement_title_re_1.match(text)
        if match:
            return Requirement(
                gs_anforderung_reference=match.group(1),
                summary=match.group(3),
                gs_absicherung=match.group(6),
                gs_verantwortliche=match.group(5),
            )
        else:
            match = cls._requirement_title_re_2.match(text)
            if match:
                return Requirement(
                    gs_anforderung_reference=match.group(1),
                    summary=match.group(3),
                    gs_absicherung=match.group(4),
                    gs_verantwortliche=match.group(6),
                )
            else:
                raise errors.ValueHttpError(
                    f"Could not parse requirement title: {text}"
                )

    @classmethod
    def _parse_requirements(cls, paragraphs):
        requirements = []
        skip_normal = True
        while paragraphs.next():
            if paragraphs.current.style.name == "Heading 3":
                skip_normal = False
                requirement = cls._parse_requirement_title(paragraphs.current.text)
                if requirement.summary == "ENTFALLEN":
                    skip_normal = True
                    continue
                cls._parse_requirement_texts(paragraphs, requirement)
                requirements.append(requirement)
            elif paragraphs.current.style.name == "Normal" and skip_normal:
                continue
            else:
                paragraphs.previous()
                break

        return requirements

    @classmethod
    def _parse_requirement_categories(cls, paragraphs):
        requirements = []
        while paragraphs.next():
            if (
                paragraphs.current.style.name == "Heading 2"
                and paragraphs.current.text
                in (
                    "Basis-Anforderungen",
                    "Standard-Anforderungen",
                    "Anforderungen bei erhöhtem Schutzbedarf",
                )
            ):
                requirements.extend(cls._parse_requirements(paragraphs))

            elif paragraphs.current.style.name == "Heading 1":
                paragraphs.previous()
                break

        return requirements

    @classmethod
    def _parse_gs_baustein_title(cls, paragraphs):
        while paragraphs.next():
            if paragraphs.current.style.name == "Title":
                match = cls._gs_baustein_title_re.match(paragraphs.current.text)
                if match:
                    return GSBaustein(reference=match.group(1), title=match.group(3))
                else:
                    raise errors.ValueHttpError(
                        f"Could not parse GS baustein title: {paragraphs.current.text}"
                    )

    @classmethod
    def _parse_requirements_from_baustein(cls, paragraphs):
        while paragraphs.next():
            if (
                paragraphs.current.style.name == "Heading 1"
                and paragraphs.current.text == "Anforderungen"
            ):
                return cls._parse_requirement_categories(paragraphs)

    @classmethod
    def _parse_baustein(cls, paragraphs):
        gs_baustein = cls._parse_gs_baustein_title(paragraphs)
        gs_baustein.requirements = cls._parse_requirements_from_baustein(paragraphs)
        return gs_baustein


def get_word_temp_file():
    with NamedTemporaryFile(suffix=".docx") as temp_file:
        yield temp_file


router = APIRouter()


@cbv(router)
class ImportGSBausteinView:
    kwargs = RequirementsView.kwargs

    def __init__(
        self,
        projects: ProjectsView = Depends(ProjectsView),
        crud: CRUDOperations[GSBaustein] = Depends(CRUDOperations),
    ):
        self._projects = projects
        self._crud = crud

    @router.post(
        "/projects/{project_id}/requirements/gs-baustein",
        status_code=201,
        response_class=Response,
        **kwargs,
    )
    def upload_gs_baustein(
        self,
        project_id: int,
        upload_file: UploadFile,
        temp_file: NamedTemporaryFile = Depends(get_word_temp_file),
    ):
        project = self._projects.get_project(project_id)
        shutil.copyfileobj(upload_file.file, temp_file.file)

        # Parse GS Baustein and save it and its requirements in the database
        gs_baustein = GSBausteinParser.parse(temp_file.name)
        for requirement in gs_baustein.requirements:
            requirement.project = project
        self._crud.create_in_db(gs_baustein)
