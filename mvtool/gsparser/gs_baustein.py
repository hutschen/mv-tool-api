# Copyright (C) 2023 Helmar Hutschenreuter
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.


from typing import cast

import docx
from docx.opc.exceptions import PackageNotFoundError

from .common import (
    GSAnforderung,
    GSBaustein,
    GSBausteinTitle,
    GSParseError,
    parse_gs_anforderung_title,
    parse_gs_baustein_title,
)


class _ParagraphsWrapper:
    def __init__(self, paragraphs):
        self.paragraphs = paragraphs
        self.index = 0

    def reset(self):
        self.index = 0

    def next(self):
        if self.index >= len(self.paragraphs):
            return None
        else:
            self.index += 1
            return self.paragraphs[self.index - 1]

    def previous(self):
        if self.index <= 1:
            return None
        else:
            self.index -= 1
            return self.paragraphs[self.index - 1]

    @property
    def current(self):
        if self.index > len(self.paragraphs):
            return None
        else:
            return self.paragraphs[self.index - 1]


def _parse_gs_anforderung_text(paragraphs: _ParagraphsWrapper):
    while paragraphs.next():
        if paragraphs.current.style.name == "Normal":
            yield cast(str, paragraphs.current.text).strip()
        else:
            paragraphs.previous()
            break


def _parse_gs_anforderungen(paragraphs: _ParagraphsWrapper):
    while paragraphs.next():
        if paragraphs.current.style.name == "Heading 3":
            yield GSAnforderung(
                title=parse_gs_anforderung_title(paragraphs.current.text),
                text=_parse_gs_anforderung_text(paragraphs),
            )
        elif paragraphs.current.style.name == "Normal":
            continue
        else:
            paragraphs.previous()
            break


def _parse_gs_anforderungen_subsections(paragraphs: _ParagraphsWrapper):
    while paragraphs.next():
        if paragraphs.current.style.name == "Heading 2" and (
            cast(str, paragraphs.current.text).lower()
            in (
                "basis-anforderungen",
                "standard-anforderungen",
                "anforderungen bei erhöhtem schutzbedarf",
            )
        ):
            for gs_anforderung in _parse_gs_anforderungen(paragraphs):
                yield gs_anforderung

        elif paragraphs.current.style.name == "Heading 1":
            paragraphs.previous()
            break


def _parse_gs_anforderungen_section(paragraphs: _ParagraphsWrapper):
    while paragraphs.next():
        if (
            paragraphs.current.style.name == "Heading 1"
            and paragraphs.current.text == "Anforderungen"
        ):
            return _parse_gs_anforderungen_subsections(paragraphs)


def _parse_gs_baustein_title(paragraphs: _ParagraphsWrapper) -> GSBausteinTitle:
    while paragraphs.next():
        if paragraphs.current.style.name == "Title":
            return parse_gs_baustein_title(paragraphs.current.text)


def _parse_gs_baustein(paragraphs: _ParagraphsWrapper) -> GSBaustein:
    return GSBaustein(
        title=_parse_gs_baustein_title(paragraphs),
        gs_anforderungen=_parse_gs_anforderungen_section(paragraphs),
    )


def parse_gs_baustein_word_file(file_name) -> GSBaustein:
    try:
        word_document = docx.Document(file_name)
    except PackageNotFoundError as e:
        raise GSParseError("Word file seems to be corrupt") from e
    paragraphs = _ParagraphsWrapper(word_document.paragraphs)
    return _parse_gs_baustein(paragraphs)
